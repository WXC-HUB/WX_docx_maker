import requests
import os
import itchat
import re
from xml.dom.minidom import parseString
import time

import selenium.webdriver

from docx import Document
from docx.shared import Inches

from tqdm import tqdm, trange

driver = 0
doConfirm = True
docxpath = 'docx-'+time.strftime("%Y-%m-%d", time.localtime(time.time()))


@itchat.msg_register(itchat.content.SHARING)
def print_content(msg):
    global doConfirm
    if(msg['ToUserName'] == 'filehelper'):
        DOMTree = parseString(msg['Content'])
        mymsg = DOMTree.documentElement
        myappmsg = mymsg.getElementsByTagName("appmsg")[0]
        mytitle = myappmsg.getElementsByTagName('title')[0].childNodes[0].data
        myurl = myappmsg.getElementsByTagName('url')[0].childNodes[0].data
        a = input("检测到链接输入，是否以如下设置创建文件？\n（输入N来取消，输入Y将在本次运行中不再询问，回车或其他输入将继续）：\ntittle:%s\nurl:%s\n" % (
            mytitle, myurl)) if doConfirm else ''
        if(not a == 'N'):
            try:
                AnalysePages(myurl, mytitle)

            except BaseException as r:
                print('发生错误：%s' % r)

            print('结果已经保存，等待下一次输入！\n')
            if(a == 'Y'):
                doConfirm = False


def StartMonitoring():
    itchat.auto_login()
    print('登录成功，请将要下载的网页发送到文件助手当中！')
    print('提示：本工具的多线程模式没有经过调试，原理上说，支持一次性传入大量链接，但也许会存在超时、重复下载等问题')
    print('提示：本工具在下载图片和生成docx文件时默认覆盖同名目录下同名文件，请做好备份')
    itchat.run()


def initdriver():
    global driver
    driver = selenium.webdriver.PhantomJS(
        executable_path="./phantomjs-2.1.1-windows/bin/phantomjs.exe")


def taketyp(mystr):
    return (re.findall(r"\?wx_fmt=(.+?)$", mystr)[0])


def AnalysePages(url, title):
    driver.get(url)
    data = driver.find_elements_by_class_name("rich_pages")
    if(len(data) < 4):
        print('rich_pages模式检测到的图片数少于4张，开始尝试图片检测\n注意：图片模式可能同时插入较多无关图片！')
        data2 = driver.find_elements_by_tag_name("img")
        if(len(data2) < 4):
            print('图片匹配模式检测到的图片数少于4张，已终止本次分析\n提示：请检查链接的格式，如果图片较少，建议手动建立word文档')
            return
        else:
            data = data2
    print('开始下载图片资源')
    ii = 0
    for i in tqdm(data, ascii=True):
        _url = i.get_attribute('data-src')
        if(not _url):
            continue
        r = requests.get(_url)
        if(taketyp(_url) == 'png' or taketyp(_url) == 'jpg' or taketyp(_url) == 'jpeg'):
            #print('下载：'+ _url+'\r', end='')
            if(not os.path.exists(title)):
                os.mkdir(title)
            with open(title+'/pic'+str(ii)+'.'+taketyp(_url), 'wb+') as f:
                f.write(r.content)
            ii = ii+1

    print('项目名：'+title+'资源下载结束！开始创建对应的docx文件，文件将被保存到'+docxpath+'文件夹内')
    MakeDOCX(title, title)


def takenum(mystr):
    return int(re.findall(r"pic(.+?)\.", mystr)[0])


def MakeDOCX(path, title):
    if(not os.path.exists(docxpath)):
        os.mkdir(docxpath)
    document = Document()
    document.add_heading(title, 0)
    picdir = os.listdir(path)
    picdir.sort(key=takenum)
    for pic in tqdm(picdir, ascii=True):
        #print('正在插入图片'+pic+'\r', end='')
        document.add_picture(path+'/'+pic, Inches(5.8))
    document.save(docxpath+'/'+title+'.docx')
    print('文件创建完成!'+title+'.docx')


if __name__ == '__main__':
    try:
        initdriver()
        StartMonitoring()
    except BaseException as a:
        print(a)
    input()
   